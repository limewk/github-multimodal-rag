"""
文件上传与内容解析 API
支持 PDF / Word (.docx) / Markdown / 纯文本 / 图片
解析后的文本可直接作为智能问答的上下文
"""
from __future__ import annotations

import base64
import io
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel

router = APIRouter(prefix="/upload", tags=["upload"])

MAX_FILE_SIZE  = 15 * 1024 * 1024   # 15 MB
MAX_TEXT_CHARS = 25_000              # 向 LLM 传递的最大字符数

TEXT_EXTENSIONS  = {".pdf", ".docx", ".doc", ".md", ".markdown", ".txt", ".rst", ".csv"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff"}
ALLOWED_EXTENSIONS = TEXT_EXTENSIONS | IMAGE_EXTENSIONS

IMAGE_MIME = {
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif":  "image/gif",
    ".webp": "image/webp",
    ".bmp":  "image/bmp",
    ".tiff": "image/tiff",
}


# ─────────────────────────────────────────────────────────────────
# Response model
# ─────────────────────────────────────────────────────────────────
class UploadResponse(BaseModel):
    filename:         str
    ext:              str
    content_type:     str             # "text" | "image"
    text:             Optional[str]   = None
    image_base64:     Optional[str]   = None
    image_media_type: Optional[str]   = None
    char_count:       int             = 0
    page_count:       int             = 0
    truncated:        bool            = False


# ─────────────────────────────────────────────────────────────────
# Endpoint
# ─────────────────────────────────────────────────────────────────
@router.post("", response_model=UploadResponse, include_in_schema=False)
@router.post("/", response_model=UploadResponse, summary="上传并解析文件")
async def upload_file(file: UploadFile = File(...)):
    """
    上传文件并提取内容：
    - PDF / Word / Markdown / TXT / CSV → 提取纯文本
    - PNG / JPG / GIF / WebP / BMP     → Base64 编码（供视觉模型分析）
    """
    filename = file.filename or "unnamed"
    ext      = Path(filename).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"不支持的文件类型「{ext or '无扩展名'}」。"
                   f"支持：PDF · Word · Markdown · TXT · CSV · PNG/JPG/GIF/WebP",
        )

    raw = await file.read()
    if len(raw) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"文件超过大小限制（最大 {MAX_FILE_SIZE // 1024 // 1024} MB）",
        )

    # ── 图片 ─────────────────────────────────────────────────────
    if ext in IMAGE_EXTENSIONS:
        return UploadResponse(
            filename=filename, ext=ext.lstrip("."),
            content_type="image",
            image_base64=base64.b64encode(raw).decode(),
            image_media_type=IMAGE_MIME.get(ext, "image/jpeg"),
            char_count=len(raw),
        )

    # ── 文本文件 ──────────────────────────────────────────────────
    text, pages = _extract_text(raw, ext)
    truncated   = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS] + "\n\n[……内容已截断，仅显示前 25 000 字符……]"

    return UploadResponse(
        filename=filename, ext=ext.lstrip("."),
        content_type="text",
        text=text,
        char_count=len(text),
        page_count=pages,
        truncated=truncated,
    )


# ─────────────────────────────────────────────────────────────────
# Text extractors
# ─────────────────────────────────────────────────────────────────
def _extract_text(raw: bytes, ext: str) -> tuple[str, int]:
    if ext == ".pdf":
        return _from_pdf(raw)
    if ext in (".docx", ".doc"):
        return _from_docx(raw)
    if ext == ".csv":
        return _from_csv(raw)
    # Markdown / TXT / RST – plain UTF-8
    try:
        text = raw.decode("utf-8", errors="replace")
    except Exception:
        text = raw.decode("latin-1", errors="replace")
    return text, 1


def _from_pdf(raw: bytes) -> tuple[str, int]:
    # Try pypdf first, fall back to PyPDF2
    for mod_name, reader_attr in [("pypdf", "PdfReader"), ("PyPDF2", "PdfReader")]:
        try:
            import importlib
            mod = importlib.import_module(mod_name)
            reader = getattr(mod, reader_attr)(io.BytesIO(raw))
            pages  = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(p for p in pages if p.strip()), len(pages)
        except ImportError:
            continue
        except Exception as e:
            return f"[PDF 解析错误：{e}]", 0

    return (
        "[PDF 解析不可用：请执行 pip install pypdf]\n\n"
        "安装后重新上传即可解析 PDF 内容。",
        0,
    )


def _from_docx(raw: bytes) -> tuple[str, int]:
    try:
        import docx  # python-docx
        doc   = docx.Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_txt:
                    parts.append(row_txt)
        return "\n\n".join(parts), len(doc.paragraphs)
    except ImportError:
        return (
            "[Word 解析不可用：请执行 pip install python-docx]\n\n"
            "安装后重新上传即可解析 Word 文档内容。",
            0,
        )
    except Exception as e:
        return f"[Word 解析错误：{e}]", 0


def _from_csv(raw: bytes) -> tuple[str, int]:
    try:
        import csv
        text = raw.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows   = [" | ".join(row) for row in reader]
        return "\n".join(rows), len(rows)
    except Exception as e:
        return f"[CSV 解析错误：{e}]", 0
